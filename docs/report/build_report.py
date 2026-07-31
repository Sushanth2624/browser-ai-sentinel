#!/usr/bin/env python3
"""Builds the Browser AI Sentinel capstone final report as a .docx, matching the RACE/REVA
format confirmed against capstone 1's own mentor-reviewed, submitted report (read-only
reference — nothing from that repo is copied or modified; only the format specs were extracted:
A4 page, 1in margins, Times New Roman throughout, Normal 12pt/1.5 spacing, Heading 1 14pt bold,
Heading 2/3 12pt bold, Caption 12pt bold).

This is a fresh, standalone generator — same mechanism as capstone 1's build script (chapter/
figure/table auto-numbering via small counters), no code shared between the two repos.

Staged build (see /root/.claude/plans — this file grows chapter by chapter across sessions):
  Stage 1 (done): infrastructure + full skeleton + front matter.
  Stage 2+: chapter content, filled in incrementally — search "STAGE" markers below.

Run: .venv/bin/python build_report.py
"""
from __future__ import annotations

import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor

TIMES = "Times New Roman"
OUT_PATH = "Browser_AI_Sentinel_Final_Report_Sushanth_Sridhar.docx"

# ---------------------------------------------------------------------------
# Document + style setup
# ---------------------------------------------------------------------------

doc = Document()


def setup_page():
    section = doc.sections[0]
    section.page_width = Inches(8.268)   # A4
    section.page_height = Inches(11.693)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)


def setup_styles():
    normal = doc.styles["Normal"]
    normal.font.name = TIMES
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(8)

    h1 = doc.styles["Heading 1"]
    h1.font.name = TIMES
    h1.font.size = Pt(14)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.line_spacing = 1.5
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(14)

    h2 = doc.styles["Heading 2"]
    h2.font.name = TIMES
    h2.font.size = Pt(12)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.line_spacing = 1.5
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after = Pt(8)

    # "Heading 3 Custom" in the reference doc — python-docx can't rename Heading 3 cleanly
    # across all Word versions, so add a genuinely custom style instead.
    if "Heading 3 Custom" not in [s.name for s in doc.styles]:
        from docx.enum.style import WD_STYLE_TYPE
        h3 = doc.styles.add_style("Heading 3 Custom", WD_STYLE_TYPE.PARAGRAPH)
        h3.base_style = doc.styles["Heading 3"]
    else:
        h3 = doc.styles["Heading 3 Custom"]
    h3.font.name = TIMES
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0, 0, 0)
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(6)

    if "Caption" not in [s.name for s in doc.styles]:
        from docx.enum.style import WD_STYLE_TYPE
        cap = doc.styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        cap = doc.styles["Caption"]
    cap.font.name = TIMES
    cap.font.size = Pt(12)
    cap.font.bold = True
    cap.font.italic = False
    cap.font.color.rgb = RGBColor(0, 0, 0)
    cap.paragraph_format.line_spacing = 1.0
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(10)


setup_page()
setup_styles()

# ---------------------------------------------------------------------------
# Chapter / figure / table numbering state
# ---------------------------------------------------------------------------

_state = {"chapter": 0, "fig": 0, "table": 0, "sub": 0, "subsub": 0}
_figures_list: list[tuple[str, str]] = []
_tables_list: list[tuple[str, str]] = []


def set_chapter(n: int):
    _state["chapter"] = n
    _state["fig"] = 0
    _state["table"] = 0
    _state["sub"] = 0
    _state["subsub"] = 0


def next_fig() -> str:
    _state["fig"] += 1
    return f"{_state['chapter']}.{_state['fig']}"


def next_table() -> str:
    _state["table"] += 1
    return f"{_state['chapter']}.{_state['table']}"


# ---------------------------------------------------------------------------
# Paragraph-level helpers
# ---------------------------------------------------------------------------

def para(text="", bold=False, italic=False, align=None, size=None, space_after=None, style=None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.font.name = TIMES
    if size:
        run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if align:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def placeholder(text: str):
    """Marks genuinely-unknown information plainly, rather than leaving it silently blank —
    same pattern the reference report used ([Designation], [Year], etc.)."""
    return para(text, italic=True)


def page_break():
    doc.add_page_break()


def chapter_heading(text: str, number: int):
    set_chapter(number)
    page_break()
    p = doc.add_paragraph(style="Heading 1")
    run = p.add_run(f"Chapter {number}: {text}")
    run.font.name = TIMES
    run.font.size = Pt(14)
    run.bold = True
    return p


def front_heading(text: str):
    """Front-matter section heading (Abstract, Acknowledgment, ...) — Heading 1 style, no
    chapter numbering."""
    page_break()
    p = doc.add_paragraph(style="Heading 1")
    run = p.add_run(text)
    run.font.name = TIMES
    run.font.size = Pt(14)
    run.bold = True
    return p


def subheading(text: str):
    _state["sub"] += 1
    _state["subsub"] = 0
    label = f"{_state['chapter']}.{_state['sub']}"
    p = doc.add_paragraph(style="Heading 2")
    run = p.add_run(f"{label} {text}")
    run.font.name = TIMES
    run.font.size = Pt(12)
    run.bold = True
    return p


def subsubheading(text: str):
    _state["subsub"] += 1
    label = f"{_state['chapter']}.{_state['sub']}.{_state['subsub']}"
    p = doc.add_paragraph(style="Heading 3 Custom")
    run = p.add_run(f"{label} {text}")
    run.font.name = TIMES
    run.font.size = Pt(12)
    run.bold = True
    return p


def figure(path: str, caption_text: str, width: float = 6.0):
    label = next_fig()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(width))
    cap = doc.add_paragraph(style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.add_run(f"Fig {label}: {caption_text}")
    _figures_list.append((label, caption_text))
    return label


def table(headers: list[str], rows: list[list[str]], caption_text: str, widths=None):
    label = next_table()
    cap = doc.add_paragraph(style="Caption")
    cap.add_run(f"Table {label}: {caption_text}")
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        r = hdr_cells[i].paragraphs[0].add_run(h)
        r.font.name = TIMES
        r.font.size = Pt(11)
        r.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(val))
            r.font.name = TIMES
            r.font.size = Pt(11)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    para("", size=1)  # small breathing room after the table
    _tables_list.append((label, caption_text))
    return label


def add_rich_run(paragraph, text: str, bold=False, base_size=12):
    """Splits text on `backtick` segments and renders those spans in a monospace font instead
    of leaving literal backtick characters in the output — a genuine formatting bug found by
    actually rendering the PDF and looking at it, not assumed correct just because the script
    ran without error."""
    parts = text.split("`")
    for i, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        if i % 2 == 1:  # odd-indexed parts were inside backticks
            run.font.name = "Consolas"
            run.font.size = Pt(base_size - 1)
        else:
            run.font.name = TIMES
            run.font.size = Pt(base_size)
        run.bold = bold


def rich_para(text: str, space_after: int | None = None):
    """Paragraph-level equivalent of add_rich_run — for prose that contains `backtick`
    code/identifier spans, so they render in monospace instead of as literal backtick
    characters (the same class of bug fixed in Chapter 5.3, now handled generically)."""
    p = doc.add_paragraph()
    add_rich_run(p, text)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def code(text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    return p


def bullet(text: str, style="List Bullet"):
    p = doc.add_paragraph(text, style=style)
    for r in p.runs:
        r.font.name = TIMES
        r.font.size = Pt(12)
    return p


# ---------------------------------------------------------------------------
# Fields: page numbers, TOC
# ---------------------------------------------------------------------------

def _field(paragraph, instr_text):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instr_text
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)
    return run


def setup_footer():
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _field(p, "PAGE")


def add_toc_field():
    p = doc.add_paragraph()
    _field(p, 'TOC \\o "1-3" \\h \\z \\u')
    note = para('Right-click here and choose "Update Field" (or press F9) to generate the '
                'Table of Contents with page numbers.', italic=True, size=10)
    return p


setup_footer()

# ---------------------------------------------------------------------------
# STAGE 1 — Front matter
# ---------------------------------------------------------------------------

PROJECT_TITLE = ("Client-Side Detection of Indirect Prompt Injection and Unauthorized Data "
                  "Exfiltration in Browser-to-AI Interactions Using Multi-Indicator Content "
                  "Analysis")

# --- Cover page ---
para("A Project Report on", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
p = para(PROJECT_TITLE, bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para("Submitted in partial fulfilment for award of degree of", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para("Master of Technology", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
para("In Cybersecurity", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para("Submitted by", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para("Sushanth Sridhar", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
para("SRN: R24TF007  |  Batch: CS14", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para("Under the Guidance of", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para("Sandeep", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
placeholder("[Designation]")
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
para("", space_after=18)
para("REVA Academy for Corporate Excellence", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
para("REVA University", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
para("Rukmini Knowledge Park, Kattigenahalli,", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
para("Yelahanka, Bangalore – 560064", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para("July 2026", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

# --- Candidate's Declaration ---
front_heading("Candidate’s Declaration")
para(
    f"I, Sushanth Sridhar, hereby declare that I have completed the project work towards "
    f"Master of Technology in Cybersecurity at REVA University on the topic entitled "
    f"“{PROJECT_TITLE}” under the supervision of Sandeep, [Designation]. This report "
    f"embodies the original work done by me in partial fulfilment of the requirements for the "
    f"award of the degree for the academic year [Year]."
)
para("Place: Bengaluru                         Name of the Student: Sushanth Sridhar")
para("Date:                                        Signature of Student")

# --- Acknowledgment of Project Ownership and Usage Rights ---
front_heading("Acknowledgment of Project Ownership and Usage Rights")
para(
    "I, Sushanth Sridhar, a student enrolled in the Master of Technology in Cybersecurity "
    "Program and Batch CS14 at RACE, hereby acknowledge that any project, including but not "
    "limited to software, hardware, research, or other intellectual property created by me "
    "during my academic tenure at RACE, is the property of RACE, REVA University."
)
para(
    "I understand and agree that RACE has the exclusive rights to use, reproduce, modify, or "
    "distribute the aforementioned projects for academic, research, and further development "
    "purposes. This includes the right to monetize, commercialize, or otherwise exploit the "
    "projects as deemed fit by RACE."
)
para(
    "I acknowledge that I have no objection to RACE, REVA University, using, reproducing, or "
    "further developing the projects for the benefit of the institution and its academic "
    "community. I further affirm that any commercial or research activities related to the "
    "projects conducted by RACE shall not require additional consent or approval from me."
)
para(
    "This acknowledgment is made willingly and without any reservations. I am grateful for the "
    "education and opportunities provided by RACE, and I recognize the importance of "
    "contributing to the academic and research goals of the institution."
)
para("Place: Bengaluru                         Name of the Student: Sushanth Sridhar")
para("Date:                                        Signature of Student")

# --- Certificate ---
front_heading("Certificate")
para(
    f"This is to certify that the project work entitled “{PROJECT_TITLE}” has been "
    f"carried out by Sushanth Sridhar with SRN R24TF007, who is a bonafide student of REVA "
    f"University, submitting the second-year project report in fulfilment for the award of "
    f"Master of Technology in Cybersecurity during the academic year [Year]. The project report "
    f"has been tested for plagiarism and has passed the plagiarism test with a similarity score "
    f"of less than 15%. The project report has been approved as it satisfies the academic "
    f"requirements in respect of the project work prescribed for the said degree."
)
para("Signature of the Guide                              Signature of the Director")
para("Sandeep                    [Name of the Director]")
para("Guide                                             Director")

p = doc.add_paragraph(style="Heading 2")
r = p.add_run("External Viva")
r.font.name = TIMES
r.font.size = Pt(12)
r.bold = True
para("Names of the Examiners")
placeholder("[Name]  [Designation]  [Signature]")
placeholder("[Name]  [Designation]  [Signature]")
para("Place: Bengaluru")
para("Date:")

# --- Acknowledgment ---
front_heading("Acknowledgment")
para(
    "I would like to express my sincere gratitude to my project guide, Sandeep, for the "
    "continuous guidance, technical feedback, and encouragement provided throughout this "
    "capstone project. I am also thankful to the faculty and program office of REVA Academy for "
    "Corporate Excellence (RACE) for their support during the course of this MTech in "
    "Cybersecurity program, batch CS14."
)
para(
    "I extend my gratitude to my classmates and peers for their valuable discussions and "
    "feedback during the development and evaluation of this project. I am also thankful to my "
    "family and friends for their patience and encouragement throughout this journey."
)
para(
    "I gratefully acknowledge the support provided by the Hon’ble Chancellor, Dr. P Shayma "
    "Raju, Hon’ble Vice Chancellor, Dr. Sanjay Chitnis, and Registrar, Dr. M. Dhanamjaya, "
    "REVA University, whose vision for research-driven, industry-relevant postgraduate education "
    "made this capstone project possible."
)
para("Place: Bengaluru")
para("Date:")

# --- Similarity Index Report ---
front_heading("Similarity Index Report")
para(
    f"This is to certify that this project report titled “{PROJECT_TITLE}” was "
    f"scanned for similarity detection. The process and outcome are given below. The plagiarism "
    f"report is attached in the appendix."
)
bullet("Software Used: Turnitin")
placeholder("Date of Report Generation: [to be filled after running Turnitin]")
placeholder("Similarity Index in %: [to be filled — must be < 15%]")
placeholder("Total word count: [auto-filled by Word once final]")
para("Name of the Guide: Sandeep")
para("Place: Bengaluru                         Name of the Student: Sushanth Sridhar")
para("Date:                                        Signature of Student")
para("Verified by:")
para("Signature")
placeholder("[Verifying officer name and title to be added]")

# --- Lists ---
front_heading("List of Abbreviations")
_ABBR = [
    ("DOM", "Document Object Model"),
    ("TLS", "Transport Layer Security"),
    ("SNI", "Server Name Indication"),
    ("JA3 / JA4", "TLS client fingerprinting methods (Salesforce JA3; FoxIO JA4)"),
    ("MV3", "Chrome Extension Manifest Version 3"),
    ("CDP", "Chrome DevTools Protocol"),
    ("DLP", "Data Loss Prevention"),
    ("PII", "Personally Identifiable Information"),
    ("EDR", "Endpoint Detection and Response"),
    ("MITRE ATLAS", "Adversarial Threat Landscape for Artificial-Intelligence Systems"),
    ("OWASP", "Open Worldwide Application Security Project"),
    ("F1", "Harmonic mean of precision and recall"),
    ("TP / FP / TN / FN", "True Positive / False Positive / True Negative / False Negative"),
]
abbr_table = doc.add_table(rows=1, cols=2)
abbr_table.style = "Table Grid"
abbr_table.rows[0].cells[0].text = "Abbreviation"
abbr_table.rows[0].cells[1].text = "Expansion"
for abbr, expansion in _ABBR:
    row = abbr_table.add_row()
    row.cells[0].text = abbr
    row.cells[1].text = expansion

front_heading("List of Figures")
placeholder("[Populated at final assembly, Stage 6 — see Fig labels generated throughout the report]")

front_heading("List of Tables")
placeholder("[Populated at final assembly, Stage 6 — see Table labels generated throughout the report]")

# --- Abstract ---
front_heading("Abstract")
placeholder("(Not to exceed 1 page)")
para(
    "AI-powered browser agents (ChatGPT Atlas, Perplexity Comet, Claude in Chrome) now read and "
    "act on arbitrary web content on a user's behalf, and the open web has begun filling with "
    "content specifically crafted to hijack that behaviour — indirect prompt injection "
    "hidden in off-screen CSS, HTML comments, and metadata that a human never sees but an agent's "
    "context window does. Independent testing found ChatGPT Atlas blocked only 5.8% and "
    "Perplexity Comet only 7% of such pages, against 47–53% for conventional phishing "
    "defences in ordinary browsers, and Google recorded a 32% rise in malicious indirect-"
    "injection content between November 2025 and February 2026. Existing defences are vendor-"
    "side and reactive; a proactive, client-side sentinel that inspects page content before an "
    "agent ever reads it is largely unexplored."
)
para(
    "This project builds and evaluates such a sentinel as a standalone system: a Chrome "
    "extension that scans the live DOM of any page for six independent indicators of hidden "
    "instructions (off-screen positioning, zero-width Unicode, suspicious HTML comments, hidden "
    "alt/ARIA text, JSON-LD metadata, and visible imperative language) and combines them with a "
    "noisy-OR multi-indicator score; a Go agent bridging the extension to a local Python scoring "
    "service and a Postgres store via Chrome's native messaging API; a real Zeek/Suricata network "
    "sensor performing JA3/JA4 TLS fingerprinting to identify known AI platforms and cluster "
    "unlisted “shadow AI” services; and an outbound data-loss-prevention gate that "
    "intercepts and classifies fetch/XHR bodies to known AI domains before they are sent, holding "
    "flagged content for user approval."
)
para(
    "The system was evaluated end to end — not simulated — against a 70-page labelled "
    "synthetic dataset (30 benign, 10 deliberately weak “hard-negative” pages, 30 "
    "injected) visited by a 4-container test fleet, each container its own OS user and hostname, "
    "running the real extension, agent, and scoring pipeline. The multi-indicator detector (C) "
    "achieved F1 0.979 (precision 0.983, recall 0.975), against F1 0.911 for a visibility-only "
    "baseline (B) and F1 0.784 for a keyword-only baseline (A) — mirroring the same C > B > A "
    "pattern from the author's prior capstone on network-based C2 detection. The result that most "
    "directly supports the multi-indicator design: on the hard-negative pages, the keyword-only "
    "baseline false-positived on 32 of 40 visits and the visibility-only baseline on 20 of 40, "
    "against 2 of 40 for the multi-indicator detector."
)
para(
    "Building the system end to end surfaced nine real, documented defects and their fixes "
    "spanning every layer — a Chrome install-path assumption that silently broke content-"
    "script injection, a native-messaging registration gap, a TLS client-fingerprint (JA3) "
    "instability caused by Chrome's GREASE mechanism that silently defeated the shadow-AI "
    "clustering rule until re-keyed to JA4, and others — each caught by checking real output "
    "against expectation rather than trusting a clean-looking log, and each recorded with root "
    "cause and fix rather than smoothed over in the final narrative. The shadow-AI clustering "
    "rule was also confirmed to have genuine practical limits: once corrected, the same "
    "fingerprint that caught the two seeded “unknown AI” test domains also pulled in "
    "sixteen unrelated domains from the browser's own background traffic, demonstrating directly "
    "why the technique is a human triage signal rather than an automated verdict."
)
para(
    "The findings support the conclusion that combining independent, cheaply-computed content "
    "indicators client-side detects indirect prompt injection substantially more reliably than "
    "any single indicator alone, while remaining fully inspectable — every flagged page "
    "shows exactly which indicators fired — and that this client-side vantage point is "
    "structurally unavailable to network-only or vendor-side defences."
)
para(
    "Keywords: Prompt Injection, Indirect Prompt Injection, Browser AI Agents, Data Loss "
    "Prevention, Shadow AI, TLS Fingerprinting, JA3, JA4, MITRE ATLAS, Multi-Indicator Detection, "
    "Chrome Extension Security"
)

# --- Table of Contents ---
front_heading("Table of Contents")
add_toc_field()

# ---------------------------------------------------------------------------
# STAGE 1 — Chapter skeleton (headings only; content filled in later stages)
# ---------------------------------------------------------------------------

STAGE_PENDING = "[Content pending — see build plan for staging]"

chapter_heading("Introduction", 1)
subheading("The Shift From a Network Perimeter to an Agentic Browser Perimeter")
para(
    "Most existing web security controls assume the browser is a passive rendering engine — a "
    "human reads a page and decides what to do next. Through late 2025 and into 2026 that "
    "assumption broke down: a new class of agentic AI browser (OpenAI's ChatGPT Atlas, "
    "Perplexity's Comet, Anthropic's Claude in Chrome) reads pages and takes actions — clicking, "
    "filling forms, sending data — on the user's behalf, with no human necessarily reading the "
    "page at all."
)
para(
    "Attackers adapted immediately. Content is now deliberately crafted to be read by an AI "
    "agent rather than a human — instructions hidden in off-screen CSS, HTML comments, "
    "zero-width Unicode, and page metadata, invisible on screen but present in the DOM an agent's "
    "context window ingests. This technique is called indirect prompt injection, and it is not "
    "theoretical: Google recorded a 32% rise in malicious indirect-injection content between "
    "November 2025 and February 2026 [2]; independent red-team testing found ChatGPT Atlas "
    "blocked only 5.8% and Perplexity Comet only 7% of malicious pages designed to hijack an "
    "agent, against 47% and 53% for conventional phishing defences in ordinary Chrome and Edge "
    "[6]; and OpenAI's own security leadership has stated that prompt injection may never be "
    "fully solved for browser agents [7]."
)
para(
    "This is structurally the same detection problem the author's prior capstone addressed for "
    "network-based command-and-control (C2) traffic — an attacker hiding instructions inside a "
    "channel the defender already trusts and cannot simply block — relocated to a new layer. "
    "There, the trusted channel was DNS/HTTPS traffic and the hidden payload was a beacon to a "
    "C2 server; here, the trusted channel is ordinary web content and the hidden payload is an "
    "instruction targeting an AI agent's context window instead of a network socket."
)
para(
    "The same agentic-AI adoption wave creates two adjacent, related problems inside the exact "
    "same browser-to-AI channel. First, organisations frequently do not know which AI services "
    "their own browsers are talking to — an increasingly recognised class of shadow IT called "
    "“shadow AI”. Second, employees paste sensitive data into AI chat interfaces with no "
    "governance at all: industry reporting places the browser as the point of roughly 80% of "
    "generative-AI data leaks, with employees regularly pasting customer data, credentials, and "
    "PII directly into prompts [11]. Commercial data-loss-prevention (DLP) products for exactly "
    "this problem already exist (Strac, Nightfall, Microsoft Purview) [11], but none of the "
    "literature or tooling surveyed in Chapter 2 combines platform discovery, outbound DLP, and "
    "proactive client-side injection detection in a single evaluated system."
)
subheading("Scope and Organization of This Report")
para(
    "This report documents the design, implementation, and evaluation of Browser AI Sentinel, a "
    "standalone Chrome extension and local agent system addressing all three problems above. "
    "Chapter 2 reviews the relevant literature and establishes the gap this project addresses. "
    "Chapter 3 states the problem precisely, including what is explicitly out of scope. Chapter 4 "
    "lists the study's objectives. Chapter 5 describes the phased build-and-verify methodology "
    "actually used, the A/B/C evaluation methodology, and — in the interest of the same "
    "evidence-over-narrative standard applied throughout this project — the real problems "
    "encountered while building it and how each was found and fixed. Chapter 6 lists resource "
    "requirements. Chapters 7 and 8 cover software design and implementation. Chapter 9 covers "
    "testing and validation, and Chapter 10 presents the analysis and results, including the real "
    "precision/recall/F1 numbers from a live, non-simulated 4-endpoint test fleet. Chapter 11 "
    "closes with limitations and future scope."
)

chapter_heading("Literature Review", 2)
subheading("Indirect Prompt Injection Against AI Browser Agents")
para(
    "Indirect prompt injection occurs when a system composes untrusted external content (a web "
    "page, a document, an email) with trusted instructions in a single context window and lets a "
    "language model act on the result [1]. Brave's security research characterises this as a "
    "fundamental, currently unsolved challenge for any AI system that ingests third-party content "
    "[1]. Field observations from Zscaler ThreatLabz [5] and Palo Alto Networks' Unit 42 [4] "
    "document real campaigns hiding instructions using off-screen CSS positioning, HTML comments, "
    "and structured JSON-LD metadata that machines parse as trusted context but a human reader "
    "never sees — the exact technique categories this project's DOM scanner targets. Mozilla has "
    "separately warned of the same risk class specifically for AI coding agents [3], and Google's "
    "own telemetry showed a 32% increase in malicious indirect-injection content in a four-month "
    "window spanning late 2025 and early 2026 [2]."
)
para(
    "Independent, adversarial testing of production AI browsers found ChatGPT Atlas and "
    "Perplexity Comet blocking only 5.8% and 7% of malicious pages respectively, against 47–53% "
    "for conventional browsers defending against phishing [6]. OpenAI's own head of preparedness "
    "has stated publicly that prompt injection for browser agents “may never be fully solved” "
    "[7], and OpenAI has since shipped iterative hardening for Atlas specifically in response to "
    "red-team findings [8]. Taken together, this literature establishes indirect prompt injection "
    "as a live, growing, and — by the vendors' own admission — structurally difficult problem, "
    "not a hypothetical one."
)
subheading("Existing Detection and Red-Team Tooling")
para(
    "Existing defensive and offensive tooling for prompt injection is concentrated at two other "
    "layers, neither of which is the client-side DOM. Open-source detectors such as Rebuff [17] "
    "and Pytector operate on the prompt text itself, using heuristics, a fine-tuned classifier, or "
    "an LLM-as-judge to flag a string as adversarial before it reaches a model — useful for "
    "direct injection into an application's own input field, but not applicable to instructions "
    "hidden in a third-party web page's DOM. Offensive/red-team tools (Praetorian's Augustus [16], "
    "a Burp Suite extension for LLM injection testing, and in-browser LLM-guided fuzzing research) "
    "are built to discover injection vulnerabilities in a target system, not to protect an "
    "end-user's own browsing session. Academic benchmarks — WAInjectBench [9], PromptShield, and "
    "a 2026 study showing prompt-injection detection performance is strongly regime-dependent "
    "(a detector tuned for one deployment context does not transfer cleanly to another) [10] — "
    "evaluate detectors against curated text/image corpora, not against a live DOM as an AI agent "
    "would actually encounter it. No surveyed system runs proactively, client-side, on the "
    "rendered page itself, ahead of an agent reading it — the specific gap this project addresses."
)
subheading("Shadow AI Discovery and Enterprise DLP for AI Tools")
para(
    "“Shadow AI” — employees using AI services an organisation has not sanctioned or even "
    "identified — is now a recognised commercial security category. Vendors including Strac, "
    "Nightfall, and Microsoft Purview ship browser-extension-based DLP that scans outbound "
    "prompts and file uploads to known AI domains before they are sent, and report that the "
    "browser is where the large majority of generative-AI data leakage actually occurs [11]. "
    "These products are mature for the outbound-DLP half of the problem — this project's own DLP "
    "module (Chapter 8) is explicitly positioned as governance, not a novel research contribution, "
    "for that reason. Discovery of shadow AI itself, however, is uniformly domain-list-based in "
    "the surveyed products: an AI service is “known” only if a vendor has already catalogued its "
    "domain. None of the surveyed tools attempt network-layer fingerprint-based discovery of "
    "*unlisted* AI services, which is where this project's shadow-AI clustering module (Chapters "
    "8 and 10) differs — and, as Chapter 10 shows honestly, where a fingerprint-based approach's "
    "real practical precision limits also become visible."
)
subheading("TLS Client Fingerprinting (JA3/JA4)")
para(
    "JA3, developed at Salesforce [14], fingerprints a TLS client by hashing the cipher suites "
    "and extensions offered in its ClientHello message. JA4, from FoxIO [15], is a newer, "
    "explicitly GREASE-aware successor designed to strip out the randomised reserved values "
    "modern browsers insert specifically to prevent fingerprint ossification. The author's prior "
    "capstone used JA3/JA4 fingerprinting as one of eight behavioural indicators for detecting "
    "C2 beaconing over encrypted traffic. This project repurposes the same underlying technique "
    "for a different task — AI-platform identification and shadow-AI discovery instead of "
    "anomaly-based C2 detection — and, in the course of doing so (Chapter 5.3), independently "
    "rediscovered in a real, verified fleet-testing scenario exactly why JA4's GREASE-awareness "
    "matters: JA3 alone proved unreliable as a stable client identifier for real Chrome traffic."
)
subheading("MITRE ATLAS and OWASP LLM Top 10")
para(
    "MITRE ATLAS (Adversarial Threat Landscape for Artificial-Intelligence Systems) [12] is a "
    "living, ATT&CK-modelled knowledge base of adversary tactics and techniques specifically "
    "against AI-enabled systems, covering 16 tactics and over 80 techniques as of its most recent "
    "revision. This project's injection-detection module maps cleanly to ATLAS technique "
    "AML.T0051.001 (LLM Prompt Injection: Indirect, under the Initial Access tactic) — verified "
    "against multiple independent sources during threat-model scoping. The outbound-DLP module's "
    "mapping to ATLAS is, honestly, unresolved: secondary sources disagreed on what technique "
    "AML.T0025 actually names, and the authoritative atlas.mitre.org site could not be fetched "
    "directly to confirm (a JavaScript single-page application). OWASP's Gen AI Security Project "
    "LLM Top 10, specifically LLM02:2025 “Sensitive Information Disclosure” [13], is used as the "
    "confirmed, safer interim citation for that module instead of an unverified ATLAS ID — a "
    "deliberate choice to avoid citing something not actually checked."
)
subheading("Consolidated Gap and This Project's Contribution")
para(
    "Across the literature surveyed, three observations converge. First, indirect prompt "
    "injection against AI browser agents is a real, growing, and — by vendors' own admission — "
    "structurally unsolved problem. Second, existing defensive tooling operates either on raw "
    "prompt text or as offensive red-team tooling, not as a proactive client-side DOM sentinel "
    "running ahead of an agent reading a page. Third, shadow-AI discovery in commercial DLP "
    "products is domain-list-based, with no surveyed system attempting network-fingerprint-based "
    "discovery of unlisted AI services. This project's contribution is a single, standalone "
    "system addressing all three gaps together — proactive DOM-level multi-indicator injection "
    "detection, JA4-based shadow-AI discovery, and an outbound DLP gate — built and evaluated end "
    "to end against real, non-simulated traffic through an actual multi-endpoint test fleet, with "
    "every claim in this report backed by a stored number, a screenshot, or a citation rather than "
    "prose alone."
)

chapter_heading("Problem Statement", 3)
para(
    "AI browser agents ingest the content of arbitrary web pages into their working context and "
    "act on it, but the current generation of such agents demonstrably fails to distinguish "
    "content intended for a human reader from content specifically crafted to instruct the agent "
    "itself. No control exists at the point where this matters most — the browser, before the "
    "agent's own (often opaque, vendor-controlled) reasoning pipeline processes the page — to "
    "flag content that reads like an instruction targeting an AI rather than a human. Separately, "
    "the same browser-to-AI channel is used to send arbitrary, ungoverned data — including "
    "personally identifiable information and secrets — to third-party AI services with no "
    "visibility or approval step, and organisations frequently cannot enumerate which AI services "
    "their own users' browsers are actually talking to."
)
para(
    "This project's threat model is scoped precisely, and deliberately excludes adjacent problems "
    "that would each need a different architecture to address correctly:"
)
bullet("In scope: indirect prompt injection — malicious instructions hidden in page content "
       "(off-screen CSS, zero-width Unicode, HTML comments, hidden alt/ARIA text, JSON-LD "
       "metadata) that a human browsing normally would not see but an AI agent reading the DOM "
       "would ingest.")
bullet("In scope: identification of which AI platform a browser is communicating with, including "
       "platforms not on any predefined list (“shadow AI”), via network-layer TLS "
       "fingerprinting rather than a static domain allowlist alone.")
bullet("In scope: unauthorized transmission of sensitive data (PII, credentials, secrets) from "
       "the browser to a known AI service, intercepted and gated before the request leaves the "
       "machine.")
bullet("Out of scope: injection via the user's own typed prompts (a different, already "
       "well-covered problem — application-level prompt filtering — and not the novel "
       "contribution here).")
bullet("Out of scope: post-hoc analysis of an AI agent's response after a successful "
       "compromise — this would require access to the agent's internal state, which is not "
       "available to a client-side browser extension.")
bullet("Out of scope: blocking an AI agent that reads a page through its own privileged channel "
       "(e.g. the Chrome DevTools Protocol or an accessibility-tree API) rather than the visible "
       "DOM — this system is explicitly an intrusion detection system (IDS), not a prevention "
       "system (IPS), for that reason, exactly mirroring the posture the author's prior capstone "
       "took with Suricata/Zeek at the network layer.")

chapter_heading("Objectives of the Study", 4)
para("The study has four objectives, each mapped directly to an implemented module (Chapter 8) "
     "and a verification step (Chapter 9):")
bullet("Objective 1: Design and implement a client-side, multi-indicator DOM scanner that "
       "detects indirect prompt injection with materially better precision and recall than any "
       "single indicator used in isolation, and evaluate this claim against a real labelled "
       "dataset rather than assert it.")
bullet("Objective 2: Design and implement network-layer AI-platform identification via real "
       "TLS SNI/JA3/JA4 fingerprinting, including a heuristic for discovering AI services not on "
       "any predefined list (“shadow AI”), and evaluate the heuristic's real behaviour "
       "against live traffic rather than only synthetic test cases.")
bullet("Objective 3: Design and implement an outbound data-loss-prevention gate that intercepts "
       "and classifies content bound for known AI services before it is sent, holding "
       "PII/secret-bearing requests for explicit user approval.")
bullet("Objective 4: Build a small, realistic multi-endpoint test environment (distinct "
       "simulated users/hosts) that exercises all three modules through the actual running "
       "system — not a simulation — and produce a governance-style dashboard summarising the "
       "results, modelled on real EDR/CASB product patterns (CrowdStrike Falcon, SentinelOne "
       "Singularity, Microsoft Purview) researched during design.")

chapter_heading("Project Methodology", 5)
subheading("Phased Build-and-Verify Methodology")
para(
    "The system was built in four sequential phases, each fully verified against real running "
    "software before the next began — the same build-verify-before-proceeding discipline used in "
    "the author's prior capstone, applied here to a very different architecture."
)
bullet("Phase 1 — single-endpoint MVP: the Chrome extension, Go agent (daemon + native-messaging "
       "shim), Python scoring service, and Postgres store, wired together and verified end to "
       "end in a real (not simulated) Chrome instance via the Chrome DevTools Protocol, with the "
       "AI-platform module still a static domain-list stub.")
bullet("Phase 2 — real network sensor: a standalone Zeek/Suricata deployment (fresh install and "
       "configuration, no code or infrastructure shared with the author's prior capstone) "
       "replacing the Phase 1 stub with real SNI/JA3/JA4 extraction and a first-cut shadow-AI "
       "clustering heuristic.")
bullet("Phase 3 — labelled dataset and multi-endpoint test fleet: a 70-page synthetic, labelled "
       "dataset for the injection-detection module, and four Docker containers — each its own "
       "OS user and hostname — running the real extension, agent, and scoring pipeline against "
       "real (not mocked) traffic, producing the evaluation data in Chapter 10.")
bullet("Phase 4 — dashboard: a governance-style web dashboard over the accumulated real data, "
       "backed by new server-side aggregate API endpoints rather than raw data dumped to the "
       "client.")
subheading("Evaluation (A/B/C Comparison) Methodology")
para(
    "The injection-detection module is evaluated using the same three-configuration comparison "
    "methodology as the author's prior capstone, applied to a different indicator set: "
    "Configuration A is a keyword-only baseline (flags a page if any visible text matches an "
    "imperative-to-AI phrase pattern); Configuration B is a visibility-only baseline (flags a "
    "page if any off-screen or zero-width-Unicode indicator is present, regardless of "
    "content); Configuration C is the full multi-indicator detector, which combines all six "
    "indicators using a noisy-OR combination (Chapter 8.1) rather than a simple threshold sum. "
    "All three configurations are computed from the same underlying indicator counts in a single "
    "request to the scoring service, so A, B, and C are never subject to timing or environment "
    "differences relative to one another — only to the real page content itself."
)
para(
    "Ground truth for the evaluation is a synthetic, labelled dataset (70 pages: 30 benign, 10 "
    "“hard-negative” pages carrying exactly one weak indicator, 30 injected pages carrying "
    "two to four indicators), generated deterministically so the same dataset can be regenerated "
    "and re-scored. Precision, recall, and F1 are computed per configuration from real "
    "confusion-matrix counts recorded by the live system during a real 4-endpoint fleet run "
    "(Chapter 9), not simulated or hand-computed."
)
subheading("Problems Encountered and How They Were Resolved")
para(
    "In keeping with the same standard applied throughout this project — real evidence, not a "
    "polished-sounding narrative — this section records every material defect found while "
    "building the system, its root cause, and its fix, in the order each was discovered. Several "
    "were only found by checking real output against expectation rather than trusting a "
    "clean-looking log or a successful build."
)
_PROBLEMS = [
    ("Chrome install path", "`--load-extension` (command-line flag) does not reliably install an "
     "unpacked Manifest V3 extension on the Chrome version used in this project — a service "
     "worker spins up, but content scripts never inject. Root cause: this flag is effectively "
     "degraded on recent Chrome; the correct path is the `Extensions.loadUnpacked` DevTools "
     "Protocol command (the same one `chrome://extensions`'s “Load unpacked” button uses "
     "internally), which also correctly honours the manifest's pinned signing key for a "
     "deterministic extension ID. Fixed by switching every automated install path to that "
     "command."),
    ("Native messaging host discovery", "A Chrome profile launched with a custom "
     "`--user-data-dir` does not discover a native-messaging-host registration placed only at "
     "the user-level path; only the system-wide path was reliably found. Fixed by registering "
     "the host at both locations."),
    ("Zeek silently drops all TLS", "The virtio network interface on the deployment VM has "
     "checksum offloading enabled, which makes Zeek see “invalid” TCP checksums on "
     "outbound packets and discard them by default — confirmed by comparing `conn.log` "
     "(populated) against `ssl.log` (empty) on the same capture. Fixed with Zeek's `-C` "
     "(ignore-checksums) flag."),
    ("Shadow-AI confidence default", "The `shadow_ai_clusters.confidence` column's schema "
     "default was mistakenly `'candidate'` instead of `'observed'`, so every single-domain "
     "sighting looked like a multi-domain cluster immediately — caught by testing against "
     "real ambient background traffic on the deployment machine, not only the constructed happy "
     "path. Fixed in the schema and via a one-time correction on the live database."),
    ("Suricata multi-interface capture", "Passing `-i ens18` on the command line silently "
     "overrides a YAML config's multi-interface `af-packet` list down to a single interface. "
     "Fixed by using `--af-packet` with no value, which tells Suricata to read the interface "
     "list from the config file."),
    ("Same-host traffic invisible to the sensor", "The mock “unknown AI” test "
     "endpoints run on the same host as the sensor; same-host traffic to a local address never "
     "transits the physical network interface, so it was invisible to a sensor watching only "
     "`ens18`. Fixed by adding a second Zeek instance and a second Suricata interface, both "
     "watching loopback."),
    ("Container native-messaging gap repeated", "The same native-messaging registration gap "
     "found earlier on the host was independently reintroduced inside the Phase 3 test-fleet "
     "containers, because the fix was not carried over into the container entrypoint script — "
     "causing the first full fleet run to silently record zero rows despite every page reporting "
     "a successful visit. Fixed the same way, plus each container's native-messaging shim needed "
     "to be told which port its own daemon was listening on, since containers share ports "
     "1-to-1 with the host under the chosen networking mode."),
    ("Extension silently skipped benign pages", "The DOM scanner only messaged the local agent "
     "when it found at least one indicator, so the dataset's 30 pure-benign pages were never "
     "scored or logged at all — leaving no true-negative data for the evaluation even after "
     "the storage layer was changed to log every score. Fixed by removing the early return so "
     "every scan reports, clean or not."),
    ("Hard-negative dataset miscalibration", "The zero-width-Unicode hard-negative test pages "
     "were generated with four zero-width characters, which the scorer's own weighting treats as "
     "a maximally strong single indicator — not the intended weak, single-occurrence test "
     "case — causing that indicator alone to saturate the score. Fixed by parameterising the "
     "generator so hard-negative pages use exactly one occurrence."),
    ("JA3 instability under Chrome's GREASE", "Shadow-AI clustering, keyed on the pair (JA3, "
     "JA4), never accumulated multi-domain evidence for real browser traffic. Root cause: "
     "Chrome's GREASE mechanism randomises reserved cipher/extension values in every ClientHello, "
     "which JA3's hashing treats as signal — confirmed empirically, with ten of twelve real "
     "Chrome connections to the same two test domains each producing a distinct JA3, while JA4 "
     "(designed to strip GREASE before hashing) stayed identical across all of them. Earlier "
     "testing had only appeared to work because it used `curl`, which does not implement GREASE. "
     "Fixed by re-keying the clustering table on JA4 alone."),
    ("Postgres ORDER BY on combined aliases", "An `ORDER BY` expression combining two "
     "`SELECT`-list aliases from correlated subqueries failed with “column does not "
     "exist,” even though each alias resolves individually. Fixed by wrapping the query in a "
     "subquery and ordering the outer query instead."),
    ("Host-level, not per-container, sensor attribution", "Network-sensor-derived events are "
     "attributed to whichever endpoint's agent is tailing the sensor logs — only the host, "
     "since the test-fleet containers do not run their own sensor — so a shadow-AI event "
     "generated by one container's browsing is recorded against the host, not that container. "
     "This is a genuine architectural limitation of packet-capture-based attribution, not a bug "
     "to be silently fixed; it is stated plainly in Chapter 10 and Chapter 11 rather than left "
     "for a reviewer to discover."),
]
for title, desc in _PROBLEMS:
    p = doc.add_paragraph()
    r = p.add_run(f"{title}. ")
    r.font.name, r.font.size, r.bold = TIMES, Pt(12), True
    add_rich_run(p, desc)
    p.paragraph_format.space_after = Pt(8)

chapter_heading("Resource Requirement Specification", 6)
subheading("Hardware Requirements")
para(
    "The complete system — extension, agent, scoring service, database, network sensor, and "
    "4-container test fleet — was built and evaluated on a single Ubuntu virtual machine: 16 "
    "vCPUs, 15 GiB RAM, and roughly 1 TB of disk (908 GB free at time of evaluation). No GPU or "
    "specialised hardware is required by any component. This single-host footprint is "
    "deliberate: the system is scoped and evaluated at the scale of a university department or "
    "small company, not a hyperscale enterprise deployment."
)
table(
    ["Component", "Minimum", "Used in this project"],
    [
        ["CPU", "4 vCPUs", "16 vCPUs"],
        ["RAM", "8 GB", "15 GiB"],
        ["Disk", "20 GB free", "~1 TB (908 GB free)"],
        ["Network", "1 real NIC (for sensor capture)", "ens18 (virtio) + loopback"],
    ],
    "Hardware requirements",
)
subheading("Software Requirements")
para("All software used is open-source or freely available, and every version below is the "
     "exact version actually used, not a general recommendation.")
table(
    ["Layer", "Software", "Version"],
    [
        ["OS", "Ubuntu Server", "24.04.4 LTS"],
        ["Extension", "TypeScript / esbuild", "5.7.2 / 0.24.0"],
        ["Extension runtime", "Google Chrome (Manifest V3)", "150.0.7871.46"],
        ["Local agent", "Go", "1.22.2"],
        ["Scoring service", "Python / FastAPI / uvicorn", "3.12.3 / 0.115.6 / 0.34.0"],
        ["Storage", "PostgreSQL", "16.14"],
        ["Network sensor", "Zeek (+ JA3/JA4 zkg packages)", "8.2.1"],
        ["Network sensor", "Suricata", "8.0.5"],
        ["Test fleet", "Docker Engine", "29.1.3"],
        ["Dashboard", "React / Vite / TypeScript", "18.3.1 / 6.0.5 / 5.7.2"],
    ],
    "Software requirements",
)
subheading("Data Requirements")
para(
    "No third-party or personally identifying real-world data was used anywhere in this project. "
    "The injection-detection evaluation uses a fully synthetic, deterministically generated "
    "70-page HTML dataset (Chapter 9). The DLP module is evaluated with synthetic PII strings "
    "constructed for testing (a fabricated email address and a Luhn-valid but non-issued card "
    "number), sent only to a nonexistent path on a real AI domain so no data is actually "
    "transmitted anywhere. The network-sensor and shadow-AI modules were evaluated against real "
    "live traffic on the deployment machine itself — the author's own machine, single-user, "
    "self-consented — which is explicitly a different, lighter-weight ethical posture than "
    "monitoring a shared or third-party network, and is stated as such."
)

chapter_heading("Software Design", 7)
subheading("System Architecture")
para(
    "The system is a single-host pipeline of five long-running components plus one passive "
    "network sensor, spanning four languages chosen for the layer each is strongest at: "
    "TypeScript for the in-browser observation surface (Manifest V3 gives no other realistic "
    "choice), Go for a small, low-overhead, always-on local daemon, Python for the scoring and "
    "classification logic where a rich ecosystem of text-processing and ML-adjacent libraries "
    "matters, and SQL/Postgres as the single source of truth every other component reads from or "
    "writes to. No component trusts another's judgement silently: the extension observes and "
    "reports, the daemon relays and persists, the scoring service decides, and the dashboard only "
    "ever reads already-decided data back out."
)
rich_para(
    "The extension cannot talk to the local agent directly over a socket — Manifest V3 confines "
    "an extension to `chrome.runtime` messaging — so a short-lived native-messaging shim "
    "(`nmhost`), spawned by Chrome itself per connection, relays each message over stdio to the "
    "persistent Go daemon's HTTP API. The daemon is the only component that talks to Postgres for "
    "writes, and the only component that talks to the Python scoring service, which keeps the "
    "extension itself free of any network or database credentials. The network sensor "
    "(Zeek + Suricata) is architecturally separate from this request path entirely: it watches "
    "the NIC passively and writes its own JSON logs, which the daemon tails on a background "
    "goroutine — the sensor has no awareness that the extension or daemon exist, and vice versa."
)
figure("assets/fig_7_1_architecture.png",
       "System architecture — component layout and the two independent capture paths "
       "(request/response via native messaging, and passive network capture via the sensor).")
subheading("Data Flow: Browser Event to Stored Verdict")
rich_para(
    "Every scored event, regardless of which module produced it, follows the same five-hop path "
    "from the page to a stored row: the content script observes something on the real page (a "
    "DOM mutation for injection scanning, a `fetch`/`XMLHttpRequest` body for the DLP gate), "
    "hands it to the background service worker, which opens a native-messaging connection that "
    "Chrome turns into a fresh `nmhost` process, which relays over stdio to the already-running "
    "Go daemon. The daemon calls the Python scoring service synchronously and blocks on its "
    "verdict before replying — this matters specifically for the DLP path, where the extension "
    "must not release an outbound request until an approval decision comes back, not merely log "
    "it after the fact."
)
para(
    "The daemon writes every verdict to Postgres unconditionally, not only the ones that cross "
    "the flag threshold — a deliberate design decision, and the fix for a real bug described in "
    "Chapter 5.3 (“Extension silently skipped benign pages”): without a complete record of "
    "benign, unflagged events, there is no true-negative population and no way to compute a real "
    "precision/recall table in Chapter 9."
)
figure("assets/fig_7_2_dataflow.png",
       "Data flow for a single browser event, from DOM/network observation through scoring to a "
       "stored verdict and the return path for a DLP approval gate.")
subheading("Module Design")
rich_para(
    "The repository is organised as one top-level directory per language/responsibility boundary, "
    "which also matches how the four objectives in Chapter 4 map onto code: injection detection "
    "and the DLP gate are almost entirely `extension/` + `ai-engine/`, AI-platform/shadow-AI "
    "identification is almost entirely `agent/` + `sensor/`, and the governance dashboard is its "
    "own `dashboard/` module reading exclusively through new server-side aggregate endpoints "
    "rather than raw tables."
)
table(
    ["Directory", "Language", "Responsibility", "Objective(s)"],
    [
        ["extension/", "TypeScript", "DOM scanning, request interception, native-messaging client", "1, 3"],
        ["agent/", "Go", "Native-messaging relay, HTTP API, sensor log tailing, Postgres writes", "1, 2, 3"],
        ["ai-engine/", "Python", "Indicator scoring, PII/DLP classification, ATLAS mapping", "1, 3"],
        ["sensor/", "Zeek/Suricata config", "Passive TLS SNI/JA3/JA4 capture", "2"],
        ["db/", "SQL", "Schema — single source of truth for all modules", "1, 2, 3, 4"],
        ["dashboard/", "React/TypeScript", "Governance-style read-only visualisation", "4"],
        ["eval/", "Python", "Labelled dataset generation and A/B/C evaluation", "1, 4"],
        ["endpoints/", "Docker/Python", "Four-container multi-endpoint test fleet", "4"],
    ],
    "Module-to-objective mapping",
)
figure("assets/fig_7_3_modules.png",
       "Module / repository dependency diagram, showing which directory depends on or drives "
       "which other directory.")

chapter_heading("Implementation", 8)
subheading("Objective 1 — Multi-Indicator Prompt-Injection Detection")
rich_para(
    "The content script (`extension/src/content-isolated/injection-scan.ts`) walks the live DOM "
    "for six indicators — off-screen CSS positioning, zero-width Unicode characters, HTML "
    "comments, `hidden`/`aria-hidden` attributes on elements carrying instruction-shaped text, "
    "JSON-LD metadata blocks, and visible imperative-to-AI language — and sends the raw counts, "
    "never a pre-computed verdict, to the Python scoring service. Keeping the counting and the "
    "scoring in separate modules across a process/language boundary is deliberate: it is what "
    "makes the A/B/C comparison in Chapter 9 possible, since Configurations A and B are computed "
    "from exactly the same indicator counts as C, not from a separately run pass."
)
rich_para(
    "The scoring service combines indicator counts with a noisy-OR formula rather than a simple "
    "weighted sum. Each indicator is treated as independent evidence; the probability that a page "
    "is benign is the product of each indicator's individual “not evidence” probability, so any "
    "single strong indicator can already push the score past the flag threshold, and two "
    "moderate indicators combine to more evidence than either alone — the actual property a "
    "“multi-indicator” detector is supposed to have. An earlier version of this function used a "
    "plain weighted average instead, which failed exactly that property: a real two-indicator "
    "sample scored below threshold as a combination despite each indicator separately being "
    "enough to flag it under the single-indicator baselines (Chapter 5.3)."
)
code(
    "def score_indicators(indicators):\n"
    "    survival = 1.0  # P(none of the indicators is evidence)\n"
    "    for name, weight in INDICATOR_WEIGHTS.items():\n"
    "        count = counts.get(name, 0)\n"
    "        if count > 0:\n"
    "            # diminishing returns, caps at 3 repeats\n"
    "            strength = weight * min(count, 3) / 3\n"
    "            survival *= (1 - strength)\n"
    "    normalized = round(1 - survival, 4)\n"
    "    return InjectionScoreResponse(\n"
    "        score=normalized,\n"
    "        flagged=normalized >= FLAG_THRESHOLD,\n"
    "    )"
)
rich_para(
    "Weights (zero-width Unicode 1.0, off-screen CSS 0.9, imperative language 0.8, "
    "`aria-hidden`/`alt` 0.7, HTML comment 0.6, JSON-LD 0.5) reflect how hard each indicator is "
    "to trigger by accident: zero-width Unicode characters essentially never occur in ordinary "
    "page content, so their presence is treated as near-certain evidence, while a single "
    "imperative sentence is weighted lower since ordinary instructional web copy can resemble it. "
    "`FLAG_THRESHOLD = 0.5` and the per-indicator weights were set by judgment before Chapter 9's "
    "labelled dataset existed, and are reported, not re-tuned to the evaluation data afterwards — "
    "tuning a threshold against the same data used to report its precision/recall would inflate "
    "the numbers artificially."
)
subheading("Objective 2 — AI-Platform and Shadow-AI Discovery")
rich_para(
    "Known-platform identification is a static domain map (`knownAIDomains` in "
    "`agent/cmd/daemon/main.go`) checked against each connection's TLS SNI. Shadow-AI discovery — "
    "finding AI-shaped traffic to domains not on that list — instead clusters connections by TLS "
    "client fingerprint on the theory that a browser or SDK talking to several different "
    "unlisted hostnames with an identical, distinctive TLS fingerprint is more likely to be one "
    "AI client library than several unrelated services. The clustering key is critical here and "
    "was wrong in an earlier version: it must be JA4, not JA3."
)
rich_para(
    "JA3 (Salesforce) and JA4 (FoxIO) both hash properties of a TLS ClientHello into a short "
    "fingerprint, but only JA4 is GREASE-aware. Chrome's GREASE mechanism deliberately randomises "
    "reserved cipher-suite and extension values in every ClientHello it sends, specifically to "
    "stop exactly this kind of fingerprint-based clustering — and JA3's hash includes those "
    "randomised values, so it changes on every connection from the same real browser. This was "
    "confirmed empirically, not assumed: ten of twelve real Chrome connections to the same two "
    "test domains each produced a distinct JA3, while JA4 (which strips GREASE values before "
    "hashing) stayed identical across all twelve. `shadow_ai_clusters` is keyed on JA4 alone, "
    "with JA3 retained only as an informational `sample_ja3` column."
)
code(
    "INSERT INTO shadow_ai_clusters\n"
    "    (ja4, sample_ja3, distinct_domains, occurrence_count)\n"
    "VALUES ($1, $2, jsonb_build_array($3::text), 1)\n"
    "ON CONFLICT (ja4) DO UPDATE SET\n"
    "    sample_ja3 = $2,\n"
    "    occurrence_count =\n"
    "        shadow_ai_clusters.occurrence_count + 1,\n"
    "    distinct_domains = CASE WHEN\n"
    "        shadow_ai_clusters.distinct_domains\n"
    "        @> jsonb_build_array($3::text)\n"
    "    THEN shadow_ai_clusters.distinct_domains\n"
    "    ELSE shadow_ai_clusters.distinct_domains\n"
    "        || jsonb_build_array($3::text) END\n"
    "RETURNING jsonb_array_length(distinct_domains)"
)
rich_para(
    "A cluster is promoted to `confidence = 'candidate'` once its distinct-domain count reaches "
    "two — a single JA4 fingerprint seen talking to two or more unlisted domains is the working "
    "definition of a shadow-AI candidate used throughout this project. This heuristic's real "
    "precision limitation (it clusters by TLS client library, not by AI-specific behaviour, so "
    "any two unlisted domains sharing an HTTP client would also cluster) is reported honestly in "
    "Chapter 10 rather than left for a reviewer to discover."
)
subheading("Objective 3 — Outbound DLP / Exfiltration Gate")
rich_para(
    "The DLP gate runs in the page's own JavaScript context (`extension/src/content-main/"
    "fetch-patch.ts`, injected at `document_start` on known-AI-domain pages) so it can patch "
    "`window.fetch` and `XMLHttpRequest` before the page's own script ever gets a reference to "
    "the originals — a background-script-only interception point would miss requests the page "
    "issues before the background worker has a chance to inject anything. Manifest V3 forbids "
    "this MAIN-world script from calling `chrome.*` APIs directly, so each intercepted request "
    "body is relayed via `window.postMessage` to an ISOLATED-world content script, which bridges "
    "it to the background worker and on to the Go daemon and Python classifier — the same "
    "five-hop path as Chapter 7.2, with one difference: this path blocks. `patchedFetch` awaits "
    "an approval decision before ever calling the real `fetch`, and throws an `AbortError` "
    "instead of sending the request if the decision comes back not-approved."
)
code(
    "window.fetch = async function patchedFetch(input, init) {\n"
    "    const bodyText = await extractBodyText(init?.body ?? null);\n"
    "    if (!bodyText.trim()) return originalFetch(input, init);\n"
    "    const approved = await requestApproval(bodyText);\n"
    "    if (!approved) {\n"
    "        throw new DOMException(\n"
    "            'Blocked: sensitive content not approved', 'AbortError');\n"
    "    }\n"
    "    return originalFetch(input, init);\n"
    "};"
)
rich_para(
    "Classification itself (`ai-engine/pii_detection/detector.py`) is deliberately simple for "
    "this project's scope: named-entity regex for email, phone, SSN, and API-key/secret "
    "patterns, plus a Luhn checksum on candidate credit-card-number matches specifically to cut "
    "false positives — an arbitrary 13-to-19-digit run is not by itself card-shaped evidence, "
    "but one that also passes Luhn is much more likely to be a real (or realistically "
    "constructed) card number. A request is held for explicit user approval if any entity type "
    "matches at all; there is no partial-trust or auto-approve tier."
)
rich_para(
    "Body extraction handles the four real `fetch`/`XHR` payload shapes a browser can send — "
    "plain string, `URLSearchParams`, `FormData` (including inspecting text-like file parts up "
    "to a size cap, so a pasted-in `.txt` or `.csv` upload is still scanned, not only form "
    "fields), and `Blob` — and caps inspected text at 200,000 bytes, since a large binary upload "
    "is not text-scannable regardless. `ArrayBuffer`/`ReadableStream` bodies are explicitly out "
    "of scope rather than silently mishandled. The approval wait fails open after five seconds on "
    "the MAIN-world side specifically to keep normal page use responsive during local "
    "development; the background worker's own longer approval-flow timeout is the fail-closed "
    "boundary for a user who is reachable but simply has not responded yet."
)
subheading("Objective 4 — Multi-Endpoint Test Fleet and Dashboard")
rich_para(
    "The dashboard (`dashboard/`, React + TypeScript) never queries raw event tables directly; "
    "every panel is backed by a purpose-built aggregate endpoint in the Go daemon "
    "(`agent/internal/store/dashboard.go`), modelled on the summary-first, drill-down-second "
    "pattern used by real EDR/CASB consoles researched during design (CrowdStrike Falcon, "
    "SentinelOne Singularity, Microsoft Purview) rather than a generic data table. The top-level "
    "KPI row, for instance, is five scalar counts computed by a single query, not five separate "
    "round trips or a client-side reduction over a full event dump:"
)
code(
    "SELECT\n"
    "    (SELECT count(*) FROM injection_alerts\n"
    "        WHERE flagged = true),\n"
    "    (SELECT count(*) FROM endpoints),\n"
    "    (SELECT count(*) FROM shadow_ai_clusters\n"
    "        WHERE confidence = 'candidate'),\n"
    "    (SELECT count(*) FROM dlp_events\n"
    "        WHERE verdict = 'flagged'),\n"
    "    (SELECT count(*) FROM dlp_events\n"
    "        WHERE verdict = 'flagged' AND approved IS NULL)"
)
rich_para(
    "The multi-endpoint test fleet (`endpoints/`) is four Docker containers, each with its own "
    "OS username and hostname, each running the real extension against a real headless Chrome "
    "instance driven over the DevTools Protocol (`endpoints/driver.py`) — not four copies of a "
    "mock client. `docker-compose.yml` uses `network_mode: host` so the fleet's native-messaging "
    "shims can reach the host's Postgres and ai-engine over loopback without widening either "
    "service's network exposure beyond `127.0.0.1`, and each container's `entrypoint.sh` "
    "registers the native-messaging host at both the user-level and system-wide discovery paths "
    "and sets a per-container `DAEMON_NM_URL` so each container's own `nmhost` process reaches "
    "its own daemon rather than another container's. This fleet is what produces the real, "
    "non-simulated confusion-matrix data reported in Chapter 9."
)

chapter_heading("Testing and Validation", 9)
subheading("Build and Type Verification")
para(STAGE_PENDING, italic=True)
subheading("Functional Test Cases (Mapped to Objectives)")
para(STAGE_PENDING, italic=True)
subheading("Headline Results (Configuration A vs B vs C)")
para(STAGE_PENDING, italic=True)

chapter_heading("Analysis and Results", 10)
subheading("Headline Comparison")
para(STAGE_PENDING, italic=True)
subheading("Why Every Single Indicator Falls Short")
para(STAGE_PENDING, italic=True)
subheading("The Shadow-AI Clustering Finding")
para(STAGE_PENDING, italic=True)
subheading("Fleet-Wide Endpoint Attribution")
para(STAGE_PENDING, italic=True)
subheading("Summary of Findings")
para(STAGE_PENDING, italic=True)

chapter_heading("Conclusions and Future Scope", 11)
subheading("Limitations")
para(STAGE_PENDING, italic=True)
subheading("Future Scope")
para(STAGE_PENDING, italic=True)

# --- Bibliography ---
front_heading("Bibliography")
para(STAGE_PENDING, italic=True)

# --- Appendix ---
front_heading("Appendix")
p = doc.add_paragraph(style="Heading 2")
r = p.add_run("Plagiarism Report")
r.font.name, r.font.size, r.bold = TIMES, Pt(12), True
placeholder("[To be attached after Turnitin scan]")

p = doc.add_paragraph(style="Heading 2")
r = p.add_run("Paper Publications in a Journal/Conference Presented/White Paper")
r.font.name, r.font.size, r.bold = TIMES, Pt(12), True
placeholder("[None at time of submission]")

p = doc.add_paragraph(style="Heading 2")
r = p.add_run("Certificate for the Conference Presentation")
r.font.name, r.font.size, r.bold = TIMES, Pt(12), True
placeholder("[Not applicable]")

p = doc.add_paragraph(style="Heading 2")
r = p.add_run("Github Link")
r.font.name, r.font.size, r.bold = TIMES, Pt(12), True
para("https://github.com/Sushanth2624/browser-ai-sentinel")

p = doc.add_paragraph(style="Heading 2")
r = p.add_run("Additional Reproduction Notes")
r.font.name, r.font.size, r.bold = TIMES, Pt(12), True
para(STAGE_PENDING, italic=True)

# ---------------------------------------------------------------------------
doc.save(OUT_PATH)
print(f"Saved {OUT_PATH} ({datetime.datetime.now().isoformat(timespec='seconds')})")
